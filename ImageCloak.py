import streamlit as st
import hashlib
import datetime
import base64
import json
import os
from PIL import Image
from io import BytesIO
from stegano import lsb
from Crypto.Cipher import AES
from Crypto.Util.Padding import pad, unpad
from web3 import Web3
from docx import Document

# ----------------------------- Blockchain Setup -----------------------------
w3 = Web3(Web3.HTTPProvider("http://127.0.0.1:8545"))
if not w3.is_connected():
    st.error("❌ Blockchain connection failed.")
    st.stop()

contract_address = Web3.to_checksum_address("0x5fbdb2315678afecb367f032d93f642f64180aa3")
with open("CryptPicRegistryABI.json") as f:
    contract_abi = json.load(f)

contract = w3.eth.contract(address=contract_address, abi=contract_abi)
w3.eth.default_account = w3.eth.accounts[0]

# ----------------------------- Encryption Utils -----------------------------
def aes_encrypt(message, key):
    try:
        cipher = AES.new(key.encode('utf-8'), AES.MODE_CBC)
        encrypted = cipher.encrypt(pad(message.encode('utf-8'), AES.block_size))
        return base64.b64encode(cipher.iv + encrypted).decode('utf-8')
    except Exception as e:
        st.error(f"Encryption error: {e}")
        return None

def aes_decrypt(encrypted_message, key):
    try:
        data = base64.b64decode(encrypted_message)
        iv = data[:16]
        cipher = AES.new(key.encode('utf-8'), AES.MODE_CBC, iv)
        return unpad(cipher.decrypt(data[16:]), AES.block_size).decode('utf-8')
    except Exception as e:
        st.error(f"Decryption error: {e}")
        return None

def hash_image(image_bytes):
    return hashlib.sha256(image_bytes).hexdigest()

# ----------------------------- Streamlit GUI -----------------------------
st.set_page_config(page_title="CryptPic", layout="wide")
st.markdown("""
    <style>
        .main { background-color: #f9f9f9; }
        .stButton > button {
            font-size: 16px;
            background-color: #2c3e50;
            color: white;
            border-radius: 10px;
            padding: 10px 20px;
        }
        .stDownloadButton > button {
            font-size: 16px;
            background-color: #27ae60;
            color: white;
            border-radius: 10px;
            padding: 10px 20px;
        }
        textarea, input, .stTextInput, .stTextArea {
            font-family: "Segoe UI", sans-serif;
        }
    </style>
""", unsafe_allow_html=True)

st.markdown("<h1 style='text-align: center; color: white; background-color: #2c3e50; padding: 20px;'>🔒 CryptPic - Professional Stego Suite</h1>", unsafe_allow_html=True)

left_col, right_col = st.columns([1, 2], gap="large")

with left_col:
    uploaded_file = st.file_uploader("📂 Upload an Image", type=["png", "jpg", "jpeg"])
    key = st.text_input("🔑 AES Key (16/24/32 chars)", type="default", help="Use a strong 16/24/32 character key.")
    message = st.text_area("📝 Message to Hide", placeholder="Optional - Only for hiding.")
    action = st.radio("Action", ["Hide Data", "Show Data"], horizontal=True)
    run = st.button("🚀 Execute")

with right_col:
    if uploaded_file:
        image = Image.open(uploaded_file)
        st.image(image, caption="Selected Image", use_column_width=True)
        image_bytes = uploaded_file.read()
        uploaded_file.seek(0)

        if run:
            if len(key) not in [16, 24, 32]:
                st.warning("Key must be 16, 24, or 32 characters.")
            else:
                if action == "Hide Data":
                    if not message.strip():
                        st.warning("Message cannot be empty.")
                    else:
                        encrypted = aes_encrypt("CryptPic Verified::" + message.strip(), key)
                        if encrypted:
                            try:
                                secret_img = lsb.hide(uploaded_file, encrypted)
                                buffer = BytesIO()
                                secret_img.save(buffer, format="PNG")
                                buffer.seek(0)

                                img_hash = hash_image(buffer.getvalue())
                                timestamp = str(datetime.datetime.utcnow())
                                tx_hash = contract.functions.storeHash(img_hash, timestamp).transact()
                                w3.eth.wait_for_transaction_receipt(tx_hash)
                                tx = w3.eth.get_transaction(tx_hash)

                                # Ensure output directories exist
                                os.makedirs("output/stego_images", exist_ok=True)
                                os.makedirs("output/records", exist_ok=True)

                                # Save stego image
                                image_filename = f"output/stego_images/stego_{datetime.datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.png"
                                with open(image_filename, "wb") as f:
                                    f.write(buffer.getvalue())

                                # Save record DOCX
                                doc = Document()
                                doc.add_heading("Blockchain Transaction", 0)
                                doc.add_paragraph(f"Transaction Hash: {tx_hash.hex()}")
                                doc.add_paragraph(f"From: {tx['from']}")
                                doc.add_paragraph(f"To: {tx['to']}")
                                doc.add_paragraph(f"Timestamp: {timestamp}")
                                doc_buffer = BytesIO()
                                doc.save(doc_buffer)
                                doc_buffer.seek(0)

                                doc_filename = f"output/records/tx_record_{datetime.datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
                                with open(doc_filename, "wb") as f:
                                    f.write(doc_buffer.getvalue())

                                st.success("✅ Data hidden & registered.")
                                st.download_button("📄 Download DOCX", doc_buffer, "transaction_details.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                                st.download_button("🖼️ Download Stego Image", buffer, "stego_image.png", mime="image/png")
                            except Exception as e:
                                st.error(f"Failed to hide data: {e}")

                elif action == "Show Data":
                    try:
                        encrypted_message = lsb.reveal(uploaded_file)
                        if not encrypted_message:
                            st.warning("No hidden data found.")
                        else:
                            decrypted = aes_decrypt(encrypted_message, key)
                            if decrypted and decrypted.startswith("CryptPic Verified::"):
                                real_message = decrypted.replace("CryptPic Verified::", "")
                                img_hash = hash_image(image_bytes)
                                try:
                                    timestamp, creator = contract.functions.verifyHash(img_hash).call()
                                    st.success(f"✅ Verified\nCreator: {creator}\nTime: {timestamp}")
                                except:
                                    st.warning("⚠️ Not found on blockchain.")
                                st.text_area("📤 Retrieved Message", real_message, height=200)
                            else:
                                st.warning("⚠️ This image is NOT verified.")
                                st.text_area("📤 Retrieved Message", decrypted or "", height=200)
                    except Exception as e:
                        st.error(f"Failed to retrieve data: {e}")
